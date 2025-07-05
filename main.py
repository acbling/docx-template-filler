import os
import openpyxl
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

# ====== 字体格式设置函数 ======

def set_font_fangsong(run):
    run.font.name = '仿宋_GB2312'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    run.font.size = Pt(10.5)

def format_excel_date(excel_date):
    if not excel_date:
        return ""
    if isinstance(excel_date, datetime):
        return excel_date.strftime('%Y年%m月%d日')
    try:
        return datetime.fromordinal(int(excel_date) + 693594).strftime('%Y年%m月%d日')
    except:
        return str(excel_date)

def center_align_table_rows(table, row_indices):
    for row_idx in row_indices:
        row = table.rows[row_idx]
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def safe_fill_cell(cell, text):
    if not text:
        text = ""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.clear()
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    set_font_fangsong(run)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def safe_fill_multiline(cell, text, last_line_right_align=False, first_line_indent=False):
    if not text:
        text = ""
    for paragraph in cell.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)

    paragraphs = str(text).split('br') if 'br' in str(text) else [str(text)]
    n = len(paragraphs)

    for i, para_text in enumerate(paragraphs):
        para_text = para_text.strip()
        if not para_text:
            continue
        p = cell.add_paragraph()
        run = p.add_run(para_text)
        set_font_fangsong(run)

        if i == n - 1 and last_line_right_align:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if i == 0 and first_line_indent:
            p.paragraph_format.first_line_indent = Pt(21)

# ====== 核心文档生成函数 ======

def fill_template_preserve_formatting(excel_path, output_folder, selected_rows):
    wb = openpyxl.load_workbook(excel_path)
    ws = wb.active
    os.makedirs(output_folder, exist_ok=True)

    for row_idx in selected_rows:
        if not ws.cell(row=row_idx, column=2).value:
            continue

        data = {
            '来文单位': ws.cell(row=row_idx, column=2).value,
            '发文日期': format_excel_date(ws.cell(row=row_idx, column=3).value),
            '收文日期': format_excel_date(ws.cell(row=row_idx, column=4).value),
            '文件编号': ws.cell(row=row_idx, column=5).value,
            '文件份数': ws.cell(row=row_idx, column=6).value,
            '文件页数': ws.cell(row=row_idx, column=7).value,
            '来文类型': ws.cell(row=row_idx, column=8).value,
            '公开属性': ws.cell(row=row_idx, column=9).value,
            '缓急程度': ws.cell(row=row_idx, column=10).value,
            '来文文号': ws.cell(row=row_idx, column=11).value,
            '文件标题': ws.cell(row=row_idx, column=12).value,
            '拟办意见': ws.cell(row=row_idx, column=13).value,
            '批示意见': ws.cell(row=row_idx, column=14).value,
            '传阅意见': ws.cell(row=row_idx, column=15).value,
            '办理情况': ws.cell(row=row_idx, column=16).value,
            '督办时间': format_excel_date(ws.cell(row=row_idx, column=17).value)
        }

        doc = Document("template.docx")  # 模板固定路径
        table = doc.tables[0]

        safe_fill_cell(table.cell(1, 1), data['来文单位'])
        safe_fill_cell(table.cell(1, 3), data['发文日期'])
        safe_fill_cell(table.cell(1, 5), data['收文日期'])

        safe_fill_cell(table.cell(2, 1), data['文件编号'])
        safe_fill_cell(table.cell(2, 3), data['文件份数'])
        safe_fill_cell(table.cell(2, 5), data['文件页数'])

        safe_fill_cell(table.cell(3, 1), data['来文类型'])
        safe_fill_cell(table.cell(3, 3), data['公开属性'])
        safe_fill_cell(table.cell(3, 5), data['缓急程度'])

        safe_fill_cell(table.cell(4, 1), data['来文文号'])
        safe_fill_cell(table.cell(4, 5), data['督办时间'])

        safe_fill_cell(table.cell(5, 1), data['文件标题'])
        safe_fill_multiline(table.cell(6, 1), data['拟办意见'])

        center_align_table_rows(table, [0, 1, 2, 3])

        # 输出文件名
        收文日期 = ws.cell(row=row_idx, column=4).value
        try:
            if isinstance(收文日期, datetime):
                收文日期_str = 收文日期.strftime('%Y%m%d')
            else:
                收文日期_str = datetime.fromordinal(int(收文日期) + 693594).strftime('%Y%m%d')
        except:
            收文日期_str = "日期未知"

        文件标题 = str(data['文件标题']) if data['文件标题'] else "无标题"
        文件标题短 = 文件标题[:30] + ('…' if len(文件标题) > 30 else '')
        filename = f"{收文日期_str}党委组织部收文处理笺（{文件标题短}）.docx"
        output_path = os.path.join(output_folder, filename)

        doc.save(output_path)
        print(f"✅ 已生成：{output_path}")

# ====== 图形界面 ======

class App:
    def __init__(self, master):
        self.master = master
        self.master.title("收文处理笺生成器")
        self.master.geometry("600x500")
        ttk.Label(master, text="📄 收文处理笺生成工具", font=("微软雅黑", 14)).pack(pady=10)

        ttk.Button(master, text="选择 Excel 文件", command=self.load_excel).pack(pady=5)
        self.list_frame = ttk.Frame(master)
        self.list_frame.pack(fill="both", expand=True)

        self.canvas = tk.Canvas(self.list_frame)
        self.scrollbar = ttk.Scrollbar(self.list_frame, orient="vertical", command=self.canvas.yview)
        self.scrollable_frame = ttk.Frame(self.canvas)

        self.scrollable_frame.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=self.scrollbar.set)

        self.canvas.pack(side="left", fill="both", expand=True)
        self.scrollbar.pack(side="right", fill="y")

        self.checkbox_vars = []
        self.entries = []
        ttk.Button(master, text="选择输出文件夹并生成", command=self.generate).pack(pady=10)

    def load_excel(self):
        path = filedialog.askopenfilename(title="选择Excel文件", filetypes=[("Excel 文件", "*.xlsx")])
        if not path:
            return
        self.excel_path = path
        wb = openpyxl.load_workbook(path)
        ws = wb.active
        self.entries.clear()
        for widget in self.scrollable_frame.winfo_children():
            widget.destroy()
        for row_idx in range(5, ws.max_row + 1):
            date = ws.cell(row=row_idx, column=4).value
            title = str(ws.cell(row=row_idx, column=12).value or "无标题")
            label = f"{format_excel_date(date)} | {title}"
            var = tk.BooleanVar()
            cb = ttk.Checkbutton(self.scrollable_frame, text=label, variable=var)
            cb.pack(anchor="w")
            self.checkbox_vars.append(var)
            self.entries.append(row_idx)

    def generate(self):
        output_dir = filedialog.askdirectory(title="选择保存目录")
        if not output_dir:
            return
        selected_rows = [self.entries[i] for i, var in enumerate(self.checkbox_vars) if var.get()]
        if not selected_rows:
            messagebox.showwarning("未选择", "请至少选择一条记录")
            return
        fill_template_preserve_formatting(self.excel_path, output_dir, selected_rows)
        messagebox.showinfo("完成", "Word 文件已生成！")

# ====== 主入口 ======

if __name__ == "__main__":
    root = tk.Tk()
    app = App(root)
    root.mainloop()
